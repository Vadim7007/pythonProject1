from docx import Document
import os
import configparser
import mysql.connector
from flask import Flask, request, make_response

# Создание Flask API приложения для принятие http запросов
app = Flask(__name__)

# Настройка подключения к бд
config = configparser.ConfigParser()
config.read("settings.ini")

connection = mysql.connector.connect(user=config["database"]["user"],
                                     password=config["database"]["password"],
                                     host=config["database"]["host"],
                                     database=config["database"]["db_name"])
cursor = connection.cursor(dictionary=True)

# Подключение файла-словаря со словами замены
dictionary = configparser.ConfigParser()
dictionary.read("dictionary.ini")


# Запрос данных этого дела из бд
def create_context(user_id, deal_id):
    result = {}
    for key in dictionary["deals"]:
        cursor.execute("SELECT " + key + " FROM wp_cases WHERE id = " + str(deal_id))
        result[dictionary["deals"][key]] = cursor.fetchall()[0].get(key)

    c = "'"
    for key in dictionary["users"]:
        cursor.execute("SELECT meta_value FROM wp_usermeta " +
                       "WHERE user_id = " + str(user_id) + " AND meta_key = " + c + key + c)
        result[dictionary["users"][key]] = cursor.fetchall()[0].get(key)

    try:
        result['Управляющий.ИО_Фамилия'] = result['last_name'] + " " + result['first_name']
    except Exception:
        pass
    return result


# Создание всех папок в директории пользователя
def create_skeleton(path_to, path_from):
    os.makedirs(path_to, exist_ok=True)

    for root, dirs, files in os.walk(path_from):
        for dir in dirs:
            os.makedirs(path_to + "/" + dir, exist_ok=True)


# Создание одного готового документа для дела в нужной папке
# Еще лучше добавить подсветку незамененных данных
def create_file(path_from, path_to, context):
    if path_from[len(path_from) - 4:] != "docx": return
    doc = Document(path_from)
    for paragraph in doc.paragraphs:
        for key in context:
            if key in paragraph.text:
                paragraph.text = context[key]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key in context:
                        if key in paragraph.text:
                            paragraph.text = paragraph.text.replace(key, str(context[key]))
    print(path_to)
    doc.save(path_to)


# Функция по созданию директорий для любой шаблонной папки
def create_type(user_id, deal_id, path_from, path_to):
    context = create_context(user_id, deal_id)

    create_skeleton(path_to, path_from)

    for root, dirs, files in os.walk(path_from):
        for dir in dirs:
            for root, dirs, files in os.walk(path_from + "/" + dir):
                for file in files:
                    create_file(path_from + "/" + dir + "/" + file,
                                path_to + "/" + dir + "/" + file,
                                context)


#  Основная функция по созданию готовой директории "конкурсное производство"
def create_kp(user_id, deal_id):
    path_from = "themes/asb/documents/Конкурсное производство"
    path_to = "uploads/documents/" + str(user_id) + "/" + str(deal_id) + "/Конкурсное производство"

    create_type(user_id, deal_id, path_from, path_to)


#  Основная функция по созданию готовой директории "Наблюдение"
def create_n(user_id, deal_id):
    path_from = "jjenserl.beget.tech/public_html/wp-content/themes/asb/documents/Наблюдение"
    path_to = "jjenserl.beget.tech/public_html/wp-content/uploads/documents/" + str(user_id) + "/" + str(deal_id) + "/Наблюдение"

    create_type(user_id, deal_id, path_from, path_to)


#  Основная функция по созданию готовой директории "Реализация"
def create_real(user_id, deal_id):
    path_from = "jjenserl.beget.tech/public_html/wp-content/themes/asb/documents/Реализация"
    path_to = "jjenserl.beget.tech/public_html/wp-content/uploads/documents/" + str(user_id) + "/" + str(deal_id) + "/Реализация"

    create_type(user_id, deal_id, path_from, path_to)


#  Основная функция по созданию готовой директории "Реструктуризация"
def create_res(user_id, deal_id):
    path_from = "jjenserl.beget.tech/public_html/wp-content/themes/asb/documents/Реструктуризация"
    path_to = "jjenserl.beget.tech/public_html/wp-content/uploads/documents/" + str(user_id) + "/" + str(deal_id) + "/Реструктуризация"

    create_type(user_id, deal_id, path_from, path_to)


# Post запрос
@app.route("/", methods=['POST'])
def post():
    params = request.values.to_dict()
    type = params.get("type")
    if type is None: return make_response(('BAD REQUEST', 400))
    type = int(type)
    if type == 0:
        try:
            create_kp(params["user_id"], params["deal_id"])
        except Exception:
            return make_response(('BAD REQUEST', 400))
    if type == 1:
        try:
            create_n(params["user_id"], params["deal_id"])
        except Exception:
            return make_response(('BAD REQUEST', 400))
    if type == 2:
        try:
            create_real(params["user_id"], params["deal_id"])
        except Exception:
            return make_response(('BAD REQUEST', 400))
    if type == 3:
        try:
            create_res(params["user_id"], params["deal_id"])
        except Exception:
            return make_response(('BAD REQUEST', 400))
    return make_response(('OK', 200))


# Запуск API Flask
if __name__ == "__main__":
    app.run(debug=False, port=80)
